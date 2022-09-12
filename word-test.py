# %config Completer.use_jedi = False
from dataclasses import dataclass
from pydoc import doc
from tkinter import font
from turtle import title, width
from docx import Document
from docx.shared import Cm
from docx.shared import Pt  #用来设置字体的大小
from docx.oxml.ns import qn #设置字体
from docx.shared import RGBColor  #设置字体的颜色
from docx.enum.text import WD_ALIGN_PARAGRAPH  #设置对其方式
import os

if not os.path.exists("./output"): os.makedirs("./output")

#全局样式
document = Document()
mystyle = document.styles["Normal"]
mystyle.font.name = '微软雅黑'
mystyle._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
mystyle.font.color.rgb=RGBColor(0,0,0)

#1级标题
para_heading=document.add_heading('',level=1)#返回1级标题段落对象，标题也相当于一个段落
para_heading.alignment=WD_ALIGN_PARAGRAPH.CENTER#设置为居中

run1 = para_heading.add_run('服务器监控指标')
run1.font.bold = True
run1.font.size = Pt(24)
# run1.font.name = '微软雅黑'
# run1._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
# run1.font.color.rgb=RGBColor(0,0,0)
# run1.style = mystyle

#设置日期
para_graph=document.add_paragraph('')#返回段落对象
para_graph.alignment=WD_ALIGN_PARAGRAPH.RIGHT#设置为右对齐

run2 = para_graph.add_run('日期:')
run2.font.size = Pt(10.5)
run2.font.name = '微软雅黑'
run2._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
run2.font.color.rgb=RGBColor(0,0,0)


#设置2级标题
para_heading=document.add_heading('',level=2)#返回2级标题段落对象，标题也相当于一个段落

run3 = para_heading.add_run('一、本周总结:')
run3.font.size = Pt(16)
run3.font.name = '微软雅黑'
run3._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
run3.font.color.rgb=RGBColor(0,0,0)
data_1 = ('1、本周监控','无','2、事件处理','无')

for i in data_1:
    para_graph=document.add_paragraph('')#返回段落对象
    run4 = para_graph.add_run(i)
    run4.font.size = Pt(10.5)
    run4.font.name = '微软雅黑'
    run4._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
    run4.font.color.rgb=RGBColor(0,0,0)


#设置2级标题
para_heading=document.add_heading('',level=2)#返回2级标题段落对象，标题也相当于一个段落

run3 = para_heading.add_run('二、监控指标:')
run3.font.size = Pt(16)
run3.font.name = '微软雅黑'
run3._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
run3.font.color.rgb=RGBColor(0,0,0)
data_1 = ('服务器CPU指标聚合','服务器内存指标聚合','服务器磁盘使用率','服务器网卡入口流量','服务器网卡出口流量')

for i in data_1:
    para_graph=document.add_paragraph('')#返回段落对象
    para_graph.alignment=WD_ALIGN_PARAGRAPH.CENTER#设置为居中
    run5 = para_graph.add_run(i)
    run5.font.size = Pt(14)
    run5.font.name = '微软雅黑'
    run5._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
    run5.font.color.rgb=RGBColor(0,0,0)
    document.add_page_break() # 添加 word 文件的分页
    # run5.add_page_break() # 添加 word 文件的分页
# para_graph_7 = document.add_paragraph('服务器CPU指标聚合')
# para_graph_8 = document.add_paragraph('服务器内存指标聚合')
# para_graph_9 = document.add_paragraph('服务器磁盘使用率')
# para_graph_10 = document.add_paragraph('服务器网卡入口流量')
# para_graph_11 =  document.add_paragraph('服务器网卡出口流量')
# para_graph_body2 = (para_graph_7,para_graph_8,para_graph_9,para_graph_10,para_graph_11)
# for para_graph_body2_style in para_graph2_title:
#     run.font.size = Pt(14)
    # para_graph_body2_style.style = run
    # doc.add_page_break()        # 添加 word 文件的分页




document.add_picture("D:\\project\\zabbix\\910.png",width=Cm(16.3), height=Cm(10.04),)

document.save("./output/服务器监控指标-.docx")
