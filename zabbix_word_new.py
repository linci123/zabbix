import requests
from turtle import title, width
from docx import Document
from docx.shared import Cm
from docx.shared import Pt  #用来设置字体的大小
from docx.oxml.ns import qn #设置字体
from docx.shared import RGBColor  #设置字体的颜色
from docx.enum.text import WD_ALIGN_PARAGRAPH  #设置对其方式

from docx import Document
import os
import datetime
import json



class Api:
    def __init__(self):
        self.session = requests.session()
        self.url = 'http://172.25.128.182/zabbix/index.php'
        self.username = "Admin"
        self.password = "zabbix"
        self.enter = "Sign in"
        self.current_date = datetime.datetime.now().strftime("%Y%m%d")
        requests.encoding = "utf8"
        response = self.session.post(self.url, files={"name":(None,self.username),"password":(None,self.password),"enter":(None,self.enter)})

    def fetchDataGraph(self, graph_id):
        data = self.session.get("http://172.25.128.182/zabbix/chart2.php?graphid={graph_id}&from=now-1h&to=now&height=201&width=1592&profileIdx=web.charts.filter".format(graph_id=graph_id),stream=True)
        if not os.path.exists("images/%s" % self.current_date): os.makedirs("images/%s" % self.current_date)
        with open("images/%s/%s.png" % (self.current_date, graph_id), "wb") as fd:
            fd.write(data.content)

    def testToGenerateWORD(self):
        #创建word
        document = Document()

        #1级标题
        para_heading=document.add_heading('',level=1)#返回1级标题段落对象，标题也相当于一个段落
        para_heading.alignment=WD_ALIGN_PARAGRAPH.CENTER#设置为居中

        run1 = para_heading.add_run('服务器监控指标')
        run1.font.bold = True
        run1.font.size = Pt(24)
        run1.font.name = '微软雅黑'
        run1._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
        run1.font.color.rgb=RGBColor(0,0,0)
        # run1.style = mystyle

        #设置日期
        para_graph=document.add_paragraph('')#返回段落对象
        para_graph.alignment=WD_ALIGN_PARAGRAPH.RIGHT#设置为右对齐

        run2 = para_graph.add_run('日期:%s') %self.current_date
        run2.font.size = Pt(10.5)
        run2.font.name = '微软雅黑'
        run2._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
        run2.font.color.rgb=RGBColor(0,0,0)

        #设置2级标题
        para_heading=document.add_heading('',level=2)#返回2级标题段落对象，标题也相当于一个段落

        run3 = para_heading.add_run('一、本周总结:')
        run3.font.size = Pt(16)
        run3.font.name = '微软雅黑'
        run3._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
        run3.font.color.rgb=RGBColor(0,0,0)
        data_1 = ('1、本周监控','无','2、事件处理','无')

        for i in data_1:
            para_graph=document.add_paragraph('')#返回段落对象
            run4 = para_graph.add_run(i)
            run4.font.size = Pt(10.5)
            run4.font.name = '微软雅黑'
            run4._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
            run4.font.color.rgb=RGBColor(0,0,0)

        #设置2级标题
        para_heading=document.add_heading('',level=2)#返回2级标题段落对象，标题也相当于一个段落

        run3 = para_heading.add_run('二、监控指标:')
        run3.font.size = Pt(16)
        run3.font.name = '微软雅黑'
        run3._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
        run3.font.color.rgb=RGBColor(0,0,0)
        data_1 = ('服务器CPU指标聚合','服务器内存指标聚合','服务器磁盘使用率','服务器网卡入口流量','服务器网卡出口流量')

        #获取图片的位置
        for i in data_1:
            para_graph=document.add_paragraph('')#返回段落对象
            para_graph.alignment=WD_ALIGN_PARAGRAPH.CENTER#设置为居中
            run5 = para_graph.add_run(i)
            run5.font.size = Pt(14)
            run5.font.name = '微软雅黑'
            run5._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
            run5.font.color.rgb=RGBColor(0,0,0)
            document.add_page_break() # 添加 word 文件的分页
        document.add_picture("D:\\project\\zabbix\\910.png",width=Cm(16.3), height=Cm(10.04))
        document.save("./output/服务器监控指标-%s.docx") % self.current_date

if __name__ ==  "__main__":
    
    pic_path = [
        {[]},
        {"name": "Disk average waiting time状态", "data":  (1938,)},
        {"name": "Disk space usage指标", "data":  (1936,)},
    ]
    api = Api()
    for v in pic_data:
        for x in v['data']:
        # if os.path.exists("images/%s/%s.png" % (current_date, x)):
        #    continue
            api.fetchDataGraph(x)
    api.testToGenerateWORD()

    